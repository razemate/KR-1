import os
import json
import csv
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from typing import Dict, Any, List
from datetime import datetime

class ExportService:
    """Enhanced export service supporting multiple formats"""
    
    def __init__(self, storage_path: str = "../storage/exports"):
        self.storage_path = storage_path
        os.makedirs(storage_path, exist_ok=True)
    
    def export_data(self, data: Any, format_type: str, filename: str = None) -> Dict[str, Any]:
        """Export data to specified format"""
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"export_{timestamp}"
            
            file_path = os.path.join(self.storage_path, f"{filename}.{format_type}")
            
            if format_type == 'csv':
                return self._export_csv(data, file_path)
            elif format_type == 'txt':
                return self._export_txt(data, file_path)
            elif format_type == 'json':
                return self._export_json(data, file_path)
            elif format_type == 'pdf':
                return self._export_pdf(data, file_path)
            elif format_type == 'xlsx':
                return self._export_xlsx(data, file_path)
            elif format_type == 'docx':
                return self._export_docx(data, file_path)
            else:
                return {"success": False, "error": f"Unsupported format: {format_type}"}
                
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _export_csv(self, data: Any, file_path: str) -> Dict[str, Any]:
        """Export to CSV format"""
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            if isinstance(data, list) and len(data) > 0:
                if isinstance(data[0], dict):
                    # List of dictionaries
                    fieldnames = data[0].keys()
                    writer = csv.DictWriter(f, fieldnames=fieldnames)
                    writer.writeheader()
                    writer.writerows(data)
                else:
                    # List of lists or simple values
                    writer = csv.writer(f)
                    if isinstance(data[0], (list, tuple)):
                        writer.writerows(data)
                    else:
                        writer.writerow(data)
            elif isinstance(data, dict):
                # Single dictionary
                writer = csv.DictWriter(f, fieldnames=data.keys())
                writer.writeheader()
                writer.writerow(data)
            else:
                # String or other data
                writer = csv.writer(f)
                writer.writerow([str(data)])
        
        return {"success": True, "file_path": file_path, "format": "csv"}
    
    def _export_txt(self, data: Any, file_path: str) -> Dict[str, Any]:
        """Export to TXT format"""
        with open(file_path, 'w', encoding='utf-8') as f:
            if isinstance(data, (dict, list)):
                f.write(json.dumps(data, indent=2, ensure_ascii=False))
            else:
                f.write(str(data))
        
        return {"success": True, "file_path": file_path, "format": "txt"}
    
    def _export_json(self, data: Any, file_path: str) -> Dict[str, Any]:
        """Export to JSON format"""
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False, default=str)
        
        return {"success": True, "file_path": file_path, "format": "json"}
    
    def _export_pdf(self, data: Any, file_path: str) -> Dict[str, Any]:
        """Export to PDF format"""
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title = Paragraph("KR-One Export Report", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Content
        if isinstance(data, dict):
            for key, value in data.items():
                para = Paragraph(f"<b>{key}:</b> {str(value)}", styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 6))
        elif isinstance(data, list):
            for i, item in enumerate(data):
                para = Paragraph(f"<b>Item {i+1}:</b> {str(item)}", styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 6))
        else:
            para = Paragraph(str(data), styles['Normal'])
            story.append(para)
        
        doc.build(story)
        return {"success": True, "file_path": file_path, "format": "pdf"}
    
    def _export_xlsx(self, data: Any, file_path: str) -> Dict[str, Any]:
        """Export to XLSX format"""
        if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
            # List of dictionaries - convert to DataFrame
            df = pd.DataFrame(data)
        elif isinstance(data, dict):
            # Single dictionary - convert to single-row DataFrame
            df = pd.DataFrame([data])
        elif isinstance(data, list):
            # List of values
            df = pd.DataFrame(data, columns=['Value'])
        else:
            # Single value
            df = pd.DataFrame([str(data)], columns=['Data'])
        
        df.to_excel(file_path, index=False, engine='openpyxl')
        return {"success": True, "file_path": file_path, "format": "xlsx"}
    
    def _export_docx(self, data: Any, file_path: str) -> Dict[str, Any]:
        """Export to DOCX format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('KR-One Export Report', 0)
        
        # Add content
        if isinstance(data, dict):
            for key, value in data.items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
        elif isinstance(data, list):
            for i, item in enumerate(data):
                p = doc.add_paragraph()
                p.add_run(f"Item {i+1}: ").bold = True
                p.add_run(str(item))
        else:
            doc.add_paragraph(str(data))
        
        doc.save(file_path)
        return {"success": True, "file_path": file_path, "format": "docx"}
    
    def list_exports(self) -> List[Dict[str, Any]]:
        """List all exported files"""
        exports = []
        try:
            for filename in os.listdir(self.storage_path):
                file_path = os.path.join(self.storage_path, filename)
                if os.path.isfile(file_path):
                    stat = os.stat(file_path)
                    exports.append({
                        "filename": filename,
                        "size": stat.st_size,
                        "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                    })
        except Exception as e:
            print(f"Error listing exports: {e}")
        
        return exports
    
    def delete_export(self, filename: str) -> bool:
        """Delete an exported file"""
        try:
            file_path = os.path.join(self.storage_path, filename)
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception as e:
            print(f"Error deleting export: {e}")
            return False

# Global instance
export_service = ExportService()